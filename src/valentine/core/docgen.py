# src/valentine/agents/docgen.py
"""
Document generation utilities for Valentine.

Generates Excel, PDF, Word, CSV, and other file types
that can be sent back to users via Telegram.
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
from dataclasses import dataclass
from pathlib import Path

from valentine.config import settings

logger = logging.getLogger(__name__)


@dataclass
class GeneratedDocument:
    """A generated document ready to be sent to the user."""
    file_path: str
    file_name: str  # human-readable name for download
    file_type: str  # "xlsx", "pdf", "docx", "csv", "html", "txt", "json"
    description: str = ""


class DocumentGenerator:
    """
    Generates documents using available Python libraries.
    Falls back gracefully when libraries aren't installed.
    """

    def __init__(self, output_dir: str | None = None):
        self.output_dir = Path(output_dir or settings.workspace_dir) / "documents"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _output_path(self, name: str, ext: str) -> str:
        """Generate a unique output path."""
        import uuid
        safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
        return str(self.output_dir / f"{safe_name}_{uuid.uuid4().hex[:6]}.{ext}")

    async def generate_excel(
        self,
        data: list[list],
        sheet_name: str = "Sheet1",
        headers: list[str] | None = None,
        file_name: str = "report",
        bold_header: bool = True,
    ) -> GeneratedDocument:
        """Generate an Excel (.xlsx) file.

        Args:
            data: 2D list of cell values
            sheet_name: Name for the worksheet
            headers: Optional column headers
            file_name: Name for the output file
            bold_header: Whether to bold the header row
        """
        try:
            import openpyxl
            from openpyxl.styles import Font
        except ImportError:
            # Fallback: generate CSV instead
            logger.warning("openpyxl not installed, falling back to CSV")
            return await self.generate_csv(data, headers=headers, file_name=file_name)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Write headers
        start_row = 1
        if headers:
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                if bold_header:
                    cell.font = Font(bold=True)
            start_row = 2

        # Write data
        for row_idx, row in enumerate(data, start_row):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Auto-adjust column widths
        for col in ws.columns:
            max_length = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

        path = self._output_path(file_name, "xlsx")
        wb.save(path)
        return GeneratedDocument(path, f"{file_name}.xlsx", "xlsx", f"Excel file with {len(data)} rows")

    async def generate_csv(
        self,
        data: list[list],
        headers: list[str] | None = None,
        file_name: str = "data",
    ) -> GeneratedDocument:
        """Generate a CSV file (always available, no dependencies)."""
        path = self._output_path(file_name, "csv")
        with open(path, "w", newline="") as f:
            writer = csv.writer(f)
            if headers:
                writer.writerow(headers)
            writer.writerows(data)
        return GeneratedDocument(path, f"{file_name}.csv", "csv", f"CSV with {len(data)} rows")

    async def generate_pdf(
        self,
        content: str,
        title: str = "",
        file_name: str = "document",
    ) -> GeneratedDocument:
        """Generate a PDF file.

        Tries reportlab first, then falls back to plain text file.
        """
        # Try reportlab
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            path = self._output_path(file_name, "pdf")
            doc = SimpleDocTemplate(path, pagesize=A4,
                                    leftMargin=20*mm, rightMargin=20*mm,
                                    topMargin=20*mm, bottomMargin=20*mm)
            styles = getSampleStyleSheet()
            story = []

            if title:
                story.append(Paragraph(title, styles['Title']))
                story.append(Spacer(1, 12))

            # Split content into paragraphs
            for para in content.split('\n\n'):
                if para.strip():
                    # Escape HTML characters
                    safe = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe.replace('\n', '<br/>'), styles['Normal']))
                    story.append(Spacer(1, 6))

            doc.build(story)
            return GeneratedDocument(path, f"{file_name}.pdf", "pdf", "PDF document")
        except ImportError:
            pass

        # Fallback: save as text
        logger.warning("reportlab not installed, saving as .txt")
        return await self.generate_text(content, file_name=file_name)

    async def generate_word(
        self,
        content: str,
        title: str = "",
        file_name: str = "document",
    ) -> GeneratedDocument:
        """Generate a Word (.docx) file."""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            if title:
                doc.add_heading(title, level=1)

            for para in content.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    for run in p.runs:
                        run.font.size = Pt(11)

            path = self._output_path(file_name, "docx")
            doc.save(path)
            return GeneratedDocument(path, f"{file_name}.docx", "docx", "Word document")
        except ImportError:
            logger.warning("python-docx not installed, saving as .txt")
            return await self.generate_text(content, file_name=file_name)

    async def generate_text(
        self,
        content: str,
        file_name: str = "output",
    ) -> GeneratedDocument:
        """Generate a plain text file (always available)."""
        path = self._output_path(file_name, "txt")
        with open(path, "w") as f:
            f.write(content)
        return GeneratedDocument(path, f"{file_name}.txt", "txt", "Text file")

    async def generate_json(
        self,
        data: dict | list,
        file_name: str = "data",
    ) -> GeneratedDocument:
        """Generate a JSON file."""
        path = self._output_path(file_name, "json")
        with open(path, "w") as f:
            json.dump(data, f, indent=2, default=str)
        return GeneratedDocument(path, f"{file_name}.json", "json", "JSON file")

    async def generate_html(
        self,
        html_content: str,
        file_name: str = "page",
    ) -> GeneratedDocument:
        """Generate an HTML file."""
        path = self._output_path(file_name, "html")
        with open(path, "w") as f:
            f.write(html_content)
        return GeneratedDocument(path, f"{file_name}.html", "html", "HTML file")
